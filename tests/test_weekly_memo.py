"""Tests for Feature 017: Weekly Platform Alignment Memo."""

import json

import pytest

from tests.conftest import HEADERS


# ---------------------------------------------------------------------------
# Platform Lead CRUD
# ---------------------------------------------------------------------------

def test_create_lead(client):
    r = client.post(
        "/leads/create",
        json={
            "name": "Matteo",
            "role": "Head of Platform Engineering",
            "focus_area": "Infrastructure, CI/CD",
            "description": "Leads platform engineering",
        },
        headers=HEADERS,
    )
    assert r.status_code == 200
    data = r.json()
    assert data["name"] == "Matteo"
    assert data["role"] == "Head of Platform Engineering"
    assert data["focus_area"] == "Infrastructure, CI/CD"
    assert data["active"] is True
    assert "id" in data


def test_create_lead_with_initiative_ids(client):
    r = client.post(
        "/leads/create",
        json={
            "name": "Mike",
            "role": "Head of AI Engineering",
            "focus_area": "AI, ML",
            "initiative_ids": ["abc-123", "def-456"],
        },
        headers=HEADERS,
    )
    assert r.status_code == 200
    data = r.json()
    assert data["initiative_ids"] == ["abc-123", "def-456"]


def test_update_lead(client):
    r = client.post(
        "/leads/create",
        json={
            "name": "Sterren",
            "role": "Head of Product",
            "focus_area": "Product",
        },
        headers=HEADERS,
    )
    lead_id = r.json()["id"]

    r2 = client.post(
        "/leads/update",
        json={
            "lead_id": lead_id,
            "role": "Head of Product & Design",
            "description": "Updated role",
        },
        headers=HEADERS,
    )
    assert r2.status_code == 200
    assert r2.json()["role"] == "Head of Product & Design"
    assert r2.json()["description"] == "Updated role"


def test_update_lead_not_found(client):
    r = client.post(
        "/leads/update",
        json={
            "lead_id": "00000000-0000-0000-0000-000000000000",
            "name": "Ghost",
        },
        headers=HEADERS,
    )
    assert r.status_code == 404


def test_list_leads(client):
    # Create two leads
    client.post(
        "/leads/create",
        json={"name": "Lead A", "role": "Role A", "focus_area": "Area A"},
        headers=HEADERS,
    )
    client.post(
        "/leads/create",
        json={"name": "Lead B", "role": "Role B", "focus_area": "Area B", "active": False},
        headers=HEADERS,
    )

    # List all
    r = client.get("/leads/list", headers=HEADERS)
    assert r.status_code == 200
    assert len(r.json()) == 2

    # List active only
    r2 = client.get("/leads/list?active_only=true", headers=HEADERS)
    assert r2.status_code == 200
    assert len(r2.json()) == 1
    assert r2.json()[0]["name"] == "Lead A"


def test_seed_leads(client):
    leads = [
        {"name": "Matteo", "role": "Head of Platform Engineering", "focus_area": "Infrastructure"},
        {"name": "Mike", "role": "Head of AI Engineering", "focus_area": "AI"},
        {"name": "Sterren", "role": "Head of Product", "focus_area": "Product"},
    ]
    r = client.post("/leads/seed", json={"leads": leads}, headers=HEADERS)
    assert r.status_code == 200
    assert len(r.json()) == 3

    # Seed again — should skip duplicates
    r2 = client.post("/leads/seed", json={"leads": leads}, headers=HEADERS)
    assert r2.status_code == 200
    assert len(r2.json()) == 0


def test_seed_leads_idempotent(client):
    leads = [{"name": "Deepak", "role": "Head of Knowledge", "focus_area": "Knowledge"}]
    client.post("/leads/seed", json={"leads": leads}, headers=HEADERS)
    r = client.post("/leads/seed", json={"leads": leads}, headers=HEADERS)
    assert r.status_code == 200
    assert len(r.json()) == 0

    # Verify only one exists
    r2 = client.get("/leads/list", headers=HEADERS)
    assert len(r2.json()) == 1


# ---------------------------------------------------------------------------
# Memo Generation
# ---------------------------------------------------------------------------

def test_generate_memo_empty(client):
    """Generate a memo with no data — should still succeed."""
    r = client.post(
        "/memos/generate",
        json={},
        headers=HEADERS,
    )
    assert r.status_code == 200
    data = r.json()
    assert data["status"] == "DRAFT"
    assert "id" in data
    assert data["strategic_objective"] is not None
    assert data["dashboard_snapshot"] is not None


def test_generate_memo_with_author(client):
    r = client.post(
        "/memos/generate",
        json={"author": "Brian"},
        headers=HEADERS,
    )
    assert r.status_code == 200
    assert r.json()["author"] == "Brian"


def test_generate_memo_with_custom_objective(client):
    r = client.post(
        "/memos/generate",
        json={"strategic_objective": "Ship v2 by Q3"},
        headers=HEADERS,
    )
    assert r.status_code == 200
    assert r.json()["strategic_objective"] == "Ship v2 by Q3"


def test_generate_memo_captures_dashboard_snapshot(client):
    """Create some commitments, then generate a memo — snapshot should reflect them."""
    # Create a commitment with priority
    client.post(
        "/commitments/open",
        json={"title": "Priority Task", "urgency": "NOW", "priority_order": 1},
        headers=HEADERS,
    )

    r = client.post("/memos/generate", json={}, headers=HEADERS)
    assert r.status_code == 200
    snapshot = r.json()["dashboard_snapshot"]
    assert "Priority Task" in snapshot.get("top_focus", [])


def test_generate_memo_with_leads(client):
    """Generate memo after seeding leads — audience should reflect seeded leads."""
    leads = [
        {"name": "Matteo", "role": "Platform Engineering", "focus_area": "Infrastructure"},
        {"name": "Mike", "role": "AI Engineering", "focus_area": "AI"},
    ]
    client.post("/leads/seed", json={"leads": leads}, headers=HEADERS)

    r = client.post("/memos/generate", json={}, headers=HEADERS)
    assert r.status_code == 200
    audience = r.json()["audience"]
    assert "Matteo" in audience
    assert "Mike" in audience


def test_generate_memo_groups_by_lead(client):
    """Verify lead_updates contain entries for each seeded lead."""
    leads = [
        {"name": "Matteo", "role": "Platform Engineering", "focus_area": "Infrastructure"},
        {"name": "Marina", "role": "Operations", "focus_area": "Operations"},
    ]
    client.post("/leads/seed", json={"leads": leads}, headers=HEADERS)

    r = client.post("/memos/generate", json={}, headers=HEADERS)
    assert r.status_code == 200
    updates = r.json()["lead_updates"]
    assert "Matteo" in updates
    assert "Marina" in updates
    assert updates["Matteo"]["role"] == "Platform Engineering"


# ---------------------------------------------------------------------------
# Memo CRUD
# ---------------------------------------------------------------------------

def test_update_memo(client):
    r = client.post("/memos/generate", json={}, headers=HEADERS)
    memo_id = r.json()["id"]

    r2 = client.post(
        "/memos/update",
        json={
            "memo_id": memo_id,
            "status": "FINALIZED",
            "strategic_objective": "Updated objective",
        },
        headers=HEADERS,
    )
    assert r2.status_code == 200
    assert r2.json()["status"] == "FINALIZED"
    assert r2.json()["strategic_objective"] == "Updated objective"


def test_update_memo_not_found(client):
    r = client.post(
        "/memos/update",
        json={
            "memo_id": "00000000-0000-0000-0000-000000000000",
            "status": "SENT",
        },
        headers=HEADERS,
    )
    assert r.status_code == 404


def test_list_memos(client):
    client.post("/memos/generate", json={}, headers=HEADERS)
    client.post("/memos/generate", json={"author": "Second"}, headers=HEADERS)

    r = client.get("/memos/list", headers=HEADERS)
    assert r.status_code == 200
    assert len(r.json()) == 2


def test_list_memos_filter_by_status(client):
    r1 = client.post("/memos/generate", json={}, headers=HEADERS)
    memo_id = r1.json()["id"]
    client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "FINALIZED"},
        headers=HEADERS,
    )

    client.post("/memos/generate", json={}, headers=HEADERS)

    # Only DRAFT
    r = client.get("/memos/list?status=DRAFT", headers=HEADERS)
    assert r.status_code == 200
    assert all(m["status"] == "DRAFT" for m in r.json())

    # Only FINALIZED
    r = client.get("/memos/list?status=FINALIZED", headers=HEADERS)
    assert r.status_code == 200
    assert all(m["status"] == "FINALIZED" for m in r.json())


def test_get_memo(client):
    r = client.post("/memos/generate", json={}, headers=HEADERS)
    memo_id = r.json()["id"]

    r2 = client.get(f"/memos/get?memo_id={memo_id}", headers=HEADERS)
    assert r2.status_code == 200
    assert r2.json()["id"] == memo_id


def test_get_memo_not_found(client):
    r = client.get(
        "/memos/get?memo_id=00000000-0000-0000-0000-000000000000",
        headers=HEADERS,
    )
    assert r.status_code == 404


# ---------------------------------------------------------------------------
# Memo Rendering
# ---------------------------------------------------------------------------

def test_render_memo(client):
    r = client.post("/memos/generate", json={"author": "Brian"}, headers=HEADERS)
    memo_id = r.json()["id"]

    r2 = client.get(f"/memos/render?memo_id={memo_id}", headers=HEADERS)
    assert r2.status_code == 200
    text = r2.text
    assert "AI Platform Weekly Leadership Memo" in text
    assert "From: Brian" in text
    assert "Strategic Objective" in text


def test_render_memo_not_found(client):
    r = client.get(
        "/memos/render?memo_id=00000000-0000-0000-0000-000000000000",
        headers=HEADERS,
    )
    assert r.status_code == 404


def test_render_memo_with_leads(client):
    leads = [
        {"name": "Matteo", "role": "Platform Engineering", "focus_area": "Infrastructure"},
    ]
    client.post("/leads/seed", json={"leads": leads}, headers=HEADERS)

    r = client.post("/memos/generate", json={"author": "Brian"}, headers=HEADERS)
    memo_id = r.json()["id"]

    r2 = client.get(f"/memos/render?memo_id={memo_id}", headers=HEADERS)
    assert r2.status_code == 200
    text = r2.text
    assert "Platform Engineering" in text
    assert "Matteo" in text


# ---------------------------------------------------------------------------
# Integration: End-to-end memo flow
# ---------------------------------------------------------------------------

def test_memo_end_to_end_flow(client):
    """Full flow: seed leads, create tasks, generate memo, finalize, render."""
    # 1. Seed leads
    leads = [
        {"name": "Matteo", "role": "Head of Platform Engineering", "focus_area": "Infrastructure, CI/CD"},
        {"name": "Mike", "role": "Head of AI Engineering", "focus_area": "AI, ML"},
        {"name": "Sterren", "role": "Head of Product", "focus_area": "Product"},
        {"name": "Marina", "role": "Head of Operations", "focus_area": "Operations"},
        {"name": "Deepak", "role": "Head of Knowledge Layer", "focus_area": "Knowledge"},
    ]
    r = client.post("/leads/seed", json={"leads": leads}, headers=HEADERS)
    assert r.status_code == 200
    assert len(r.json()) == 5

    # 2. Create some commitments
    client.post(
        "/commitments/open",
        json={"title": "Deploy new CI pipeline", "urgency": "NOW", "priority_order": 1},
        headers=HEADERS,
    )
    client.post(
        "/commitments/open",
        json={"title": "Review AI model results", "urgency": "SOON"},
        headers=HEADERS,
    )

    # 3. Generate memo
    r = client.post(
        "/memos/generate",
        json={"author": "Brian"},
        headers=HEADERS,
    )
    assert r.status_code == 200
    memo = r.json()
    assert memo["status"] == "DRAFT"
    assert memo["author"] == "Brian"
    assert len(memo["audience"]) == 5

    # 4. Finalize
    r = client.post(
        "/memos/update",
        json={"memo_id": memo["id"], "status": "FINALIZED"},
        headers=HEADERS,
    )
    assert r.status_code == 200
    assert r.json()["status"] == "FINALIZED"

    # 5. Render
    r = client.get(f"/memos/render?memo_id={memo['id']}", headers=HEADERS)
    assert r.status_code == 200
    text = r.text
    assert "AI Platform Weekly Leadership Memo" in text
    assert "Brian" in text
    assert "Deploy new CI pipeline" in text or "Strategic Objective" in text


# ---------------------------------------------------------------------------
# Memo Status Transitions (back-and-forth)
# ---------------------------------------------------------------------------

def test_status_draft_to_finalized(client):
    r = client.post("/memos/generate", json={}, headers=HEADERS)
    memo_id = r.json()["id"]

    r2 = client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "FINALIZED"},
        headers=HEADERS,
    )
    assert r2.status_code == 200
    assert r2.json()["status"] == "FINALIZED"


def test_status_finalized_back_to_draft(client):
    """User should be able to revert a finalized memo back to draft for edits."""
    r = client.post("/memos/generate", json={}, headers=HEADERS)
    memo_id = r.json()["id"]

    # Finalize
    client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "FINALIZED"},
        headers=HEADERS,
    )

    # Revert to draft
    r2 = client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "DRAFT"},
        headers=HEADERS,
    )
    assert r2.status_code == 200
    assert r2.json()["status"] == "DRAFT"


def test_status_finalized_to_sent(client):
    r = client.post("/memos/generate", json={}, headers=HEADERS)
    memo_id = r.json()["id"]

    client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "FINALIZED"},
        headers=HEADERS,
    )

    r2 = client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "SENT"},
        headers=HEADERS,
    )
    assert r2.status_code == 200
    assert r2.json()["status"] == "SENT"


def test_status_sent_back_to_draft(client):
    """User can revert a sent memo all the way back to draft."""
    r = client.post("/memos/generate", json={}, headers=HEADERS)
    memo_id = r.json()["id"]

    client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "SENT"},
        headers=HEADERS,
    )

    r2 = client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "DRAFT"},
        headers=HEADERS,
    )
    assert r2.status_code == 200
    assert r2.json()["status"] == "DRAFT"


def test_status_full_lifecycle(client):
    """DRAFT -> FINALIZED -> DRAFT (edit) -> FINALIZED -> SENT."""
    r = client.post("/memos/generate", json={"author": "Brian"}, headers=HEADERS)
    memo_id = r.json()["id"]
    assert r.json()["status"] == "DRAFT"

    # Finalize
    r2 = client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "FINALIZED"},
        headers=HEADERS,
    )
    assert r2.json()["status"] == "FINALIZED"

    # Revert to draft for edits
    r3 = client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "DRAFT", "strategic_objective": "Revised objective"},
        headers=HEADERS,
    )
    assert r3.json()["status"] == "DRAFT"
    assert r3.json()["strategic_objective"] == "Revised objective"

    # Re-finalize
    r4 = client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "FINALIZED"},
        headers=HEADERS,
    )
    assert r4.json()["status"] == "FINALIZED"

    # Send
    r5 = client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "SENT"},
        headers=HEADERS,
    )
    assert r5.json()["status"] == "SENT"


# ---------------------------------------------------------------------------
# Memo Export — Markdown
# ---------------------------------------------------------------------------

def test_export_md(client):
    r = client.post("/memos/generate", json={"author": "Brian"}, headers=HEADERS)
    memo_id = r.json()["id"]

    r2 = client.get(f"/memos/export-md?memo_id={memo_id}", headers=HEADERS)
    assert r2.status_code == 200
    text = r2.text
    # Verify markdown heading structure
    assert "# AI Platform Weekly Leadership Memo" in text
    assert "**From:** Brian" in text
    assert "## Strategic Objective" in text
    assert "**Status:** DRAFT" in text


def test_export_md_with_leads(client):
    leads = [
        {"name": "Matteo", "role": "Platform Engineering", "focus_area": "Infrastructure"},
    ]
    client.post("/leads/seed", json={"leads": leads}, headers=HEADERS)

    r = client.post("/memos/generate", json={"author": "Brian"}, headers=HEADERS)
    memo_id = r.json()["id"]

    r2 = client.get(f"/memos/export-md?memo_id={memo_id}", headers=HEADERS)
    assert r2.status_code == 200
    text = r2.text
    assert "### Platform Engineering" in text
    assert "Matteo" in text


def test_export_md_not_found(client):
    r = client.get(
        "/memos/export-md?memo_id=00000000-0000-0000-0000-000000000000",
        headers=HEADERS,
    )
    assert r.status_code == 404


def test_export_md_content_disposition(client):
    r = client.post("/memos/generate", json={}, headers=HEADERS)
    memo_id = r.json()["id"]

    r2 = client.get(f"/memos/export-md?memo_id={memo_id}", headers=HEADERS)
    assert r2.status_code == 200
    assert "attachment" in r2.headers.get("content-disposition", "")
    assert ".md" in r2.headers.get("content-disposition", "")


# ---------------------------------------------------------------------------
# Memo Export — Word (.docx)
# ---------------------------------------------------------------------------

def test_export_docx(client):
    r = client.post("/memos/generate", json={"author": "Brian"}, headers=HEADERS)
    memo_id = r.json()["id"]

    r2 = client.get(f"/memos/export-docx?memo_id={memo_id}", headers=HEADERS)
    assert r2.status_code == 200
    # Verify it's a valid docx (starts with PK zip header)
    assert r2.content[:2] == b"PK"
    assert "application/vnd.openxmlformats" in r2.headers.get("content-type", "")


def test_export_docx_content_disposition(client):
    r = client.post("/memos/generate", json={}, headers=HEADERS)
    memo_id = r.json()["id"]

    r2 = client.get(f"/memos/export-docx?memo_id={memo_id}", headers=HEADERS)
    assert r2.status_code == 200
    assert "attachment" in r2.headers.get("content-disposition", "")
    assert ".docx" in r2.headers.get("content-disposition", "")


def test_export_docx_not_found(client):
    r = client.get(
        "/memos/export-docx?memo_id=00000000-0000-0000-0000-000000000000",
        headers=HEADERS,
    )
    assert r.status_code == 404


def test_export_docx_contains_content(client):
    """Verify the exported docx contains the expected text content."""
    r = client.post("/memos/generate", json={"author": "Brian"}, headers=HEADERS)
    memo_id = r.json()["id"]

    r2 = client.get(f"/memos/export-docx?memo_id={memo_id}", headers=HEADERS)
    assert r2.status_code == 200

    # Parse the docx bytes to verify content
    import io
    from docx import Document

    doc = Document(io.BytesIO(r2.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "AI Platform Weekly Leadership Memo" in full_text
    assert "Brian" in full_text
    assert "Strategic Objective" in full_text


# ---------------------------------------------------------------------------
# Export after status transitions
# ---------------------------------------------------------------------------

def test_export_finalized_memo(client):
    """Export should work on finalized memos and reflect the status."""
    r = client.post("/memos/generate", json={"author": "Brian"}, headers=HEADERS)
    memo_id = r.json()["id"]

    client.post(
        "/memos/update",
        json={"memo_id": memo_id, "status": "FINALIZED"},
        headers=HEADERS,
    )

    # Export as md
    r2 = client.get(f"/memos/export-md?memo_id={memo_id}", headers=HEADERS)
    assert r2.status_code == 200
    assert "**Status:** FINALIZED" in r2.text

    # Export as docx
    r3 = client.get(f"/memos/export-docx?memo_id={memo_id}", headers=HEADERS)
    assert r3.status_code == 200
    assert r3.content[:2] == b"PK"

    # Verify docx contains finalized status
    import io
    from docx import Document

    doc = Document(io.BytesIO(r3.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "FINALIZED" in full_text


# ---------------------------------------------------------------------------
# Auth enforcement
# ---------------------------------------------------------------------------

def test_leads_require_auth(client):
    r = client.post(
        "/leads/create",
        json={"name": "No Auth", "role": "Test", "focus_area": "Test"},
    )
    assert r.status_code == 401


def test_memos_require_auth(client):
    r = client.post("/memos/generate", json={})
    assert r.status_code == 401


def test_export_requires_auth(client):
    r = client.get("/memos/export-md?memo_id=00000000-0000-0000-0000-000000000000")
    assert r.status_code == 401

    r2 = client.get("/memos/export-docx?memo_id=00000000-0000-0000-0000-000000000000")
    assert r2.status_code == 401
