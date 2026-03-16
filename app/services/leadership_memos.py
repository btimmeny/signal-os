"""CRUD and generation logic for leadership memos.

Feature 019: Dynamic AI Platform Weekly Leadership Memo System.
All leadership roles, ownership, and organisational references are resolved
dynamically from the ``platform_leads`` table — nothing is hard-coded.

Memo sections follow a strict order:
  1. Strategic Direction
  2. Progress This Week
  3. Why It Matters
  4. Next Platform Moves
  5. Leadership Execution

Content is rendered as leadership narrative paragraphs.
Target length: 600-800 words.
"""

from __future__ import annotations

import io
import json
import logging
import os
import subprocess
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Optional

from sqlalchemy.orm import Session

from app.models import (
    Commitment,
    CommitmentStatus,
    Initiative,
    InitiativeCommitmentLink,
    InitiativeStatus,
    LeadershipMemo,
    MemoStatus,
    PlatformLead,
    StrategicTheme,
    ThemeStatus,
)

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

DEFAULT_STRATEGIC_OBJECTIVE = (
    "The organization is building an AI-native engineering platform where "
    "intelligent agents participate directly in the software development "
    "lifecycle."
)

# Strategy memory: persistent understanding of platform direction
STRATEGY_MEMORY = (
    "The organization is building an AI-native engineering platform where "
    "intelligent agents participate directly in the software development "
    "lifecycle. This platform has four major architectural components: "
    "AI-Native SDLC, Agent Infrastructure, Knowledge Layer, and Data Platform. "
    "The AI-Native SDLC enables AI agents to assist with code creation, review, "
    "remediation, and deployment. Agent Infrastructure provides the orchestration "
    "layer for AI agents, including swarm architectures and event-driven workflows. "
    "The Knowledge Layer organizes specifications, architecture decisions, and "
    "operational knowledge so that both engineers and AI agents can use it. "
    "The Data Platform ensures that AI systems have reliable access to structured "
    "internal data and that outputs from AI systems can be distributed across "
    "the organization."
)

PLATFORM_PILLARS = [
    "AI-Native SDLC",
    "Agent Infrastructure",
    "Knowledge Layer",
    "Data Platform",
]

# Repo-relative paths for file persistence
_MEMO_DIR = Path("leadership-memos/ai-platform/weekly")
_EXPORT_DIR = _MEMO_DIR / "exports"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _week_start(dt: Optional[datetime] = None) -> datetime:
    """Return the Monday 00:00 UTC of the week containing *dt*."""
    if dt is None:
        dt = datetime.now(timezone.utc)
    monday = dt - timedelta(days=dt.weekday())
    return monday.replace(hour=0, minute=0, second=0, microsecond=0, tzinfo=timezone.utc)


def _parse_json_field(val: object, default: object = None) -> object:
    """Parse a JSON-encoded text field, returning *default* if not parseable."""
    if val is None:
        return default
    if isinstance(val, (list, dict)):
        return val
    try:
        return json.loads(val)
    except (json.JSONDecodeError, TypeError):
        return default


def _repo_root() -> Path:
    """Return the repository root directory."""
    root = os.environ.get("SIGNAL_OS_REPO_ROOT")
    if root:
        return Path(root)
    candidate = Path.cwd()
    if (candidate / "app").is_dir():
        return candidate
    return Path("/home/ubuntu/repos/signal-os")


def _active_leads(db: Session) -> list[PlatformLead]:
    """Return all active platform leads ordered by creation date."""
    return (
        db.query(PlatformLead)
        .filter(PlatformLead.active == 1)
        .order_by(PlatformLead.created_at.asc())
        .all()
    )


def _lead_initiative_ids(lead: PlatformLead) -> list[str]:
    """Parse the JSON initiative_ids on a lead, returning empty list on failure."""
    if not lead.initiative_ids:
        return []
    try:
        return json.loads(lead.initiative_ids)
    except (json.JSONDecodeError, TypeError):
        return []


# ---------------------------------------------------------------------------
# Narrative section builders
# ---------------------------------------------------------------------------


def _build_narrative_strategic_direction(
    db: Session,
    override: Optional[str] = None,
) -> str:
    """Build the *Strategic Direction* section (2-3 paragraphs).

    Describes what the organization is building, how the current week fits
    into the broader strategy, and how the platform pillars are evolving.
    Reinforces the strategy memory.
    """
    if override:
        return override

    # Paragraph 1: What we're building (from strategy memory)
    theme = (
        db.query(StrategicTheme)
        .filter(StrategicTheme.status == ThemeStatus.ACTIVE)
        .order_by(StrategicTheme.created_at.desc())
        .first()
    )
    if theme and theme.description:
        para1 = theme.description
    else:
        para1 = STRATEGY_MEMORY

    # Paragraph 2: How this week fits into broader strategy
    active_initiatives = (
        db.query(Initiative)
        .filter(Initiative.status == InitiativeStatus.ACTIVE)
        .order_by(Initiative.created_at.asc())
        .all()
    )
    active_count = len(active_initiatives)

    # Map initiatives to pillars by keyword matching
    pillar_inits: dict[str, list[str]] = {p: [] for p in PLATFORM_PILLARS}
    for init in active_initiatives:
        title_lower = init.title.lower()
        for pillar in PLATFORM_PILLARS:
            if any(kw in title_lower for kw in pillar.lower().split()):
                pillar_inits[pillar].append(init.title)
                break

    active_pillars = [p for p, inits in pillar_inits.items() if inits]
    if active_pillars:
        para2 = (
            f"This week, the platform team is advancing work across "
            f"{len(active_pillars)} of the four platform pillars: "
            f"{', '.join(active_pillars)}. "
            f"There are {active_count} active initiatives driving progress "
            f"toward the AI-native engineering platform."
        )
    else:
        para2 = (
            f"The platform team is executing across {active_count} active "
            f"initiatives this week, continuing to build toward the four "
            f"architectural pillars: {', '.join(PLATFORM_PILLARS)}."
        )

    return f"{para1}\n\n{para2}"


def _build_narrative_progress(db: Session, leads: list[PlatformLead]) -> str:
    """Build the *Progress This Week* section.

    Highlights 1-3 meaningful milestones from completed work or significant
    platform progress. Written in narrative form, focused on outcomes rather
    than activities. Includes one major platform milestone.
    """
    # Recently closed commitments (completed work)
    recently_closed = (
        db.query(Commitment)
        .filter(Commitment.status == CommitmentStatus.CLOSED)
        .order_by(Commitment.closed_at.desc())
        .limit(5)
        .all()
    )

    active_initiatives = (
        db.query(Initiative)
        .filter(Initiative.status == InitiativeStatus.ACTIVE)
        .all()
    )
    init_map = {str(i.id): i for i in active_initiatives}

    # Build milestone narratives from completed work
    milestones: list[str] = []
    for commitment in recently_closed[:3]:
        # Try to connect to an initiative/pillar
        links = (
            db.query(InitiativeCommitmentLink)
            .filter(InitiativeCommitmentLink.commitment_id == commitment.id)
            .all()
        )
        init_name = None
        for link in links:
            init = init_map.get(str(link.initiative_id))
            if init:
                init_name = init.title
                break

        if init_name:
            milestones.append(
                f"{commitment.title} was completed as part of the {init_name} "
                f"initiative, advancing the platform's capabilities."
            )
        else:
            milestones.append(
                f"{commitment.title} was completed, contributing to overall "
                f"platform execution progress."
            )

    if not milestones:
        # No recently closed items — describe active work momentum
        all_open = (
            db.query(Commitment)
            .filter(Commitment.status != CommitmentStatus.CLOSED)
            .count()
        )
        milestones.append(
            f"The platform team is actively executing across {all_open} open "
            f"commitments, building momentum toward key platform milestones."
        )

    return " ".join(milestones)


def _build_narrative_why_it_matters(db: Session) -> str:
    """Build the *Why It Matters* section.

    Explains why the progress from this week is strategically important.
    Connects progress to platform strategy. Focuses on capabilities being
    unlocked rather than tasks completed.
    """
    active_initiatives = (
        db.query(Initiative)
        .filter(Initiative.status == InitiativeStatus.ACTIVE)
        .all()
    )

    recently_closed = (
        db.query(Commitment)
        .filter(Commitment.status == CommitmentStatus.CLOSED)
        .order_by(Commitment.closed_at.desc())
        .limit(5)
        .all()
    )

    # Identify which pillars are being advanced
    pillar_activity: dict[str, int] = {p: 0 for p in PLATFORM_PILLARS}
    for init in active_initiatives:
        title_lower = init.title.lower()
        for pillar in PLATFORM_PILLARS:
            if any(kw in title_lower for kw in pillar.lower().split()):
                pillar_activity[pillar] += 1
                break

    active_pillars = [p for p, count in pillar_activity.items() if count > 0]
    closed_count = len(recently_closed)

    if active_pillars and closed_count > 0:
        text = (
            f"This week's progress is strategically significant because it "
            f"advances {len(active_pillars)} of the four platform pillars"
        )
        if len(active_pillars) <= 3:
            text += f" ({', '.join(active_pillars)})"
        text += (
            f". The completion of {closed_count} commitment(s) represents "
            f"tangible progress toward enabling AI-native development workflows "
            f"and strengthening the platform architecture."
        )
    elif active_pillars:
        text = (
            f"Active work across {', '.join(active_pillars)} is building "
            f"the foundation for AI-native development capabilities. Each "
            f"initiative contributes to the broader goal of enabling intelligent "
            f"agents to participate directly in the software development lifecycle."
        )
    else:
        text = (
            "The current execution focus is establishing the foundational "
            "capabilities that will enable AI-native development workflows, "
            "agent infrastructure, knowledge management, and data platform "
            "integration across the engineering organization."
        )

    return text


def _build_narrative_next_moves(db: Session) -> str:
    """Build the *Next Platform Moves* section.

    Describes the next stage of platform evolution, building on progress
    made this week. Forward-looking and strategic.
    """
    # Look at top-priority open commitments for forward-looking moves
    top_priority = (
        db.query(Commitment)
        .filter(
            Commitment.status != CommitmentStatus.CLOSED,
            Commitment.priority_order.isnot(None),
        )
        .order_by(Commitment.priority_order.asc())
        .limit(3)
        .all()
    )

    active_initiatives = (
        db.query(Initiative)
        .filter(Initiative.status == InitiativeStatus.ACTIVE)
        .all()
    )

    # Upcoming due items
    now = datetime.now(timezone.utc)
    cutoff = now + timedelta(days=7)
    due_soon = (
        db.query(Commitment)
        .filter(
            Commitment.status != CommitmentStatus.CLOSED,
            Commitment.due_at.isnot(None),
        )
        .all()
    )
    upcoming: list[str] = []
    for c in due_soon:
        due_at = c.due_at if c.due_at.tzinfo else c.due_at.replace(tzinfo=timezone.utc)
        if due_at <= cutoff:
            upcoming.append(c.title)

    parts: list[str] = []

    if top_priority:
        priority_titles = [c.title for c in top_priority[:3]]
        parts.append(
            f"The next stage of platform evolution focuses on "
            f"{', '.join(priority_titles)}. These represent the highest-priority "
            f"moves that will advance the platform architecture."
        )

    if active_initiatives and not top_priority:
        init_titles = [i.title for i in active_initiatives[:3]]
        parts.append(
            f"The platform will build on current momentum by advancing "
            f"{', '.join(init_titles)}, extending the AI-native engineering "
            f"capabilities across the organization."
        )

    if upcoming:
        parts.append(
            f"{len(upcoming)} deliverable(s) approach their target dates in "
            f"the coming week, requiring focused execution from the team."
        )

    if not parts:
        parts.append(
            "The platform will continue to expand its AI-native engineering "
            "capabilities, focusing on agent infrastructure maturity, "
            "knowledge layer development, and data platform integration."
        )

    return " ".join(parts)


def _build_narrative_leadership_execution(
    db: Session, leads: list[PlatformLead]
) -> str:
    """Build the *Leadership Execution* section.

    Each leader receives exactly three commitments. Each commitment includes
    the task and a short explanation of why it matters strategically.

    Format:
        Leader Name

        Task
        Why it matters: explanation
    """
    if not leads:
        return (
            "Leadership execution assignments are pending. Roles will be "
            "confirmed and commitments assigned once the team structure "
            "is finalized."
        )

    active_initiatives = (
        db.query(Initiative)
        .filter(Initiative.status == InitiativeStatus.ACTIVE)
        .all()
    )
    init_map = {str(i.id): i for i in active_initiatives}

    all_open = (
        db.query(Commitment)
        .filter(Commitment.status != CommitmentStatus.CLOSED)
        .order_by(Commitment.priority_order.asc().nullslast(), Commitment.opened_at.asc())
        .all()
    )
    open_by_id = {str(c.id): c for c in all_open}

    # Map commitments to initiatives
    all_links = db.query(InitiativeCommitmentLink).all()
    commitment_init: dict[str, str] = {}
    for link in all_links:
        cid = str(link.commitment_id)
        iid = str(link.initiative_id)
        if iid in init_map:
            commitment_init[cid] = init_map[iid].title

    # Map initiatives to pillars for "why it matters"
    init_pillar: dict[str, str] = {}
    for init in active_initiatives:
        title_lower = init.title.lower()
        for pillar in PLATFORM_PILLARS:
            if any(kw in title_lower for kw in pillar.lower().split()):
                init_pillar[str(init.id)] = pillar
                break

    sections: list[str] = []
    assigned_commitment_ids: set[str] = set()

    for lead in leads:
        lead_init_ids = _lead_initiative_ids(lead)

        # Gather commitments for this lead from their initiatives
        lead_commitments: list[Commitment] = []
        for iid in lead_init_ids:
            init_links = (
                db.query(InitiativeCommitmentLink)
                .filter(InitiativeCommitmentLink.initiative_id == iid)
                .all()
            )
            for link in init_links:
                cid = str(link.commitment_id)
                if cid in open_by_id and cid not in assigned_commitment_ids:
                    lead_commitments.append(open_by_id[cid])
                    assigned_commitment_ids.add(cid)

        # If not enough from initiatives, match by focus area keywords
        if len(lead_commitments) < 3:
            focus_keywords = [kw.strip().lower() for kw in lead.focus_area.split(",")]
            for c in all_open:
                cid = str(c.id)
                if cid in assigned_commitment_ids:
                    continue
                if any(kw in c.title.lower() for kw in focus_keywords):
                    lead_commitments.append(c)
                    assigned_commitment_ids.add(cid)
                if len(lead_commitments) >= 3:
                    break

        # Limit to exactly 3
        lead_commitments = lead_commitments[:3]

        # Build the section for this lead
        lead_lines = [f"**{lead.name}**", ""]
        for commitment in lead_commitments:
            cid = str(commitment.id)
            init_name = commitment_init.get(cid, "")
            # Generate "why it matters" based on initiative/pillar connection
            if init_name:
                why = (
                    f"Advances the {init_name} initiative, contributing to "
                    f"the platform's strategic execution."
                )
            else:
                why = (
                    f"Strengthens platform execution and supports progress "
                    f"toward the AI-native engineering architecture."
                )
            lead_lines.append(commitment.title)
            lead_lines.append(f"Why it matters: {why}")
            lead_lines.append("")

        if not lead_commitments:
            lead_lines.append(f"Focused on {lead.focus_area}")
            lead_lines.append(
                f"Why it matters: Ensures continued progress on "
                f"{lead.focus_area.lower()} capabilities."
            )
            lead_lines.append("")

        sections.append("\n".join(lead_lines))

    return "\n".join(sections)


# ---------------------------------------------------------------------------
# Dashboard snapshot (kept for backward compatibility with DB field)
# ---------------------------------------------------------------------------


def _gather_dashboard_snapshot(db: Session) -> dict:
    """Pull the current dashboard state for the memo snapshot."""
    all_open = (
        db.query(Commitment)
        .filter(Commitment.status != CommitmentStatus.CLOSED)
        .all()
    )

    now = datetime.now(timezone.utc)
    due_soon_cutoff = now + timedelta(days=7)

    top_focus: list[str] = []
    due_soon: list[str] = []
    active_workstreams: list[str] = []

    for c in all_open:
        if c.priority_order is not None:
            top_focus.append(c.title)
        if c.due_at:
            due_at = c.due_at if c.due_at.tzinfo else c.due_at.replace(tzinfo=timezone.utc)
            if due_at <= due_soon_cutoff:
                due_soon.append(f"{c.title} (due {c.due_at.strftime('%b %-d')})")

    active_initiatives = (
        db.query(Initiative)
        .filter(Initiative.status == InitiativeStatus.ACTIVE)
        .order_by(Initiative.created_at.asc())
        .all()
    )
    for init in active_initiatives:
        task_count = (
            db.query(InitiativeCommitmentLink)
            .join(Commitment, InitiativeCommitmentLink.commitment_id == Commitment.id)
            .filter(
                InitiativeCommitmentLink.initiative_id == init.id,
                Commitment.status != CommitmentStatus.CLOSED,
            )
            .count()
        )
        if task_count > 0:
            active_workstreams.append(f"{init.title} ({task_count} active tasks)")

    return {
        "top_focus": top_focus,
        "needs_decision": [],
        "due_soon": due_soon,
        "active_workstreams": active_workstreams,
    }


def _build_lead_updates(db: Session) -> dict:
    """Group initiatives and tasks by platform lead (JSON stored in DB)."""
    leads = _active_leads(db)

    active_initiatives = (
        db.query(Initiative)
        .filter(Initiative.status == InitiativeStatus.ACTIVE)
        .all()
    )
    init_map = {str(i.id): i for i in active_initiatives}

    all_open = (
        db.query(Commitment)
        .filter(Commitment.status != CommitmentStatus.CLOSED)
        .all()
    )
    open_by_id = {str(c.id): c for c in all_open}

    init_tasks: dict[str, list[str]] = {}
    for init in active_initiatives:
        links = (
            db.query(InitiativeCommitmentLink)
            .filter(InitiativeCommitmentLink.initiative_id == init.id)
            .all()
        )
        tasks = [open_by_id[str(lnk.commitment_id)].title
                 for lnk in links if str(lnk.commitment_id) in open_by_id]
        if tasks:
            init_tasks[str(init.id)] = tasks

    updates: dict[str, dict] = {}
    for lead in leads:
        lead_init_ids = _lead_initiative_ids(lead)

        progress: list[str] = []
        for iid in lead_init_ids:
            init = init_map.get(iid)
            if not init:
                continue
            tasks = init_tasks.get(iid, [])
            if tasks:
                for t in tasks:
                    progress.append(f"{init.title}: {t}")
            else:
                progress.append(f"{init.title}: No active tasks")

        if not lead_init_ids:
            focus_keywords = [kw.strip().lower() for kw in lead.focus_area.split(",")]
            for init in active_initiatives:
                if any(kw in init.title.lower() for kw in focus_keywords):
                    tasks = init_tasks.get(str(init.id), [])
                    for t in tasks:
                        progress.append(f"{init.title}: {t}")

        updates[lead.name] = {
            "role": lead.role,
            "focus": lead.focus_area,
            "progress": progress,
            "next_focus": [],
        }

    return updates


# ---------------------------------------------------------------------------
# Memo generation
# ---------------------------------------------------------------------------


def generate_memo(
    db: Session,
    *,
    author: Optional[str] = None,
    strategic_objective: Optional[str] = None,
) -> LeadershipMemo:
    """Generate a weekly leadership memo from current data store state.

    All leadership roles are resolved dynamically from the platform_leads
    table.  Content is stored in the existing LeadershipMemo columns as
    JSON where needed.

    Column mapping to new sections:
      strategic_objective  -> Strategic Direction
      progress_summary     -> Progress This Week
      focus_next_week      -> Why It Matters (JSON string)
      success_criteria     -> Next Platform Moves (JSON string)
      lead_updates         -> Leadership Execution (JSON object)
    """
    leads = _active_leads(db)
    snapshot = _gather_dashboard_snapshot(db)
    lead_updates = _build_lead_updates(db)

    # Build narrative sections (new 5-section structure)
    direction_text = _build_narrative_strategic_direction(db, override=strategic_objective)
    progress_text = _build_narrative_progress(db, leads)
    why_text = _build_narrative_why_it_matters(db)
    next_moves_text = _build_narrative_next_moves(db)
    execution_text = _build_narrative_leadership_execution(db, leads)

    # Audience = all active leads
    audience = [ld.name for ld in leads] if leads else []

    memo = LeadershipMemo(
        id=uuid.uuid4(),
        week_start_date=_week_start(),
        author=author,
        status=MemoStatus.DRAFT,
        strategic_objective=direction_text,
        current_priorities=json.dumps([]),
        progress_summary=progress_text,
        focus_next_week=json.dumps(why_text),
        success_criteria=json.dumps(next_moves_text),
        lead_updates=json.dumps(lead_updates),
        dashboard_snapshot=json.dumps(snapshot),
        audience=json.dumps(audience),
    )
    # Store leadership execution text in lead_updates alongside legacy data
    lead_updates["_execution_text"] = execution_text
    memo.lead_updates = json.dumps(lead_updates)

    db.add(memo)
    db.commit()
    db.refresh(memo)
    return memo


# ---------------------------------------------------------------------------
# Markdown formatting — narrative style
# ---------------------------------------------------------------------------


def format_memo_markdown(memo: LeadershipMemo) -> str:
    """Render a memo as a Markdown document using narrative paragraphs.

    Section order:
      1. Strategic Direction
      2. Progress This Week
      3. Why It Matters
      4. Next Platform Moves
      5. Leadership Execution

    Target 600-800 words.
    """
    audience = _parse_json_field(memo.audience, [])
    lead_updates = _parse_json_field(memo.lead_updates, {})

    # Unpack narrative fields (stored as plain text or JSON strings)
    progress_text = memo.progress_summary or ""
    why_text = _parse_json_field(memo.focus_next_week, "")
    if isinstance(why_text, list):
        why_text = ", ".join(why_text)
    next_moves_text = _parse_json_field(memo.success_criteria, "")
    if isinstance(next_moves_text, list):
        next_moves_text = ", ".join(next_moves_text)

    # Extract leadership execution text from lead_updates
    execution_text = ""
    if isinstance(lead_updates, dict):
        execution_text = lead_updates.pop("_execution_text", "")
    if not execution_text:
        # Fallback: build from lead_updates data
        exec_sentences: list[str] = []
        for name, info in lead_updates.items():
            if not isinstance(info, dict):
                continue
            role = info.get("role", "")
            focus = info.get("focus", "")
            exec_sentences.append(f"{name}, {role}, is focused on {focus}")
        execution_text = ". ".join(exec_sentences) + "." if exec_sentences else ""

    date_str = memo.week_start_date.strftime("%B %-d, %Y")
    status_val = memo.status.value if hasattr(memo.status, "value") else memo.status

    lines: list[str] = [
        "# AI Platform Weekly Leadership Memo",
        "",
        f"**Week Of:** {date_str}",
        f"**From:** {memo.author or 'Leadership'}",
        f"**Status:** {status_val}",
        "",
        "## Strategic Direction",
        "",
        memo.strategic_objective or DEFAULT_STRATEGIC_OBJECTIVE,
        "",
        "## Progress This Week",
        "",
        progress_text,
        "",
        "## Why It Matters",
        "",
        why_text if isinstance(why_text, str) else str(why_text),
        "",
        "## Next Platform Moves",
        "",
        next_moves_text if isinstance(next_moves_text, str) else str(next_moves_text),
        "",
        "## Leadership Execution",
        "",
        execution_text if isinstance(execution_text, str) else str(execution_text),
    ]

    return "\n".join(lines)


def format_memo_text(db: Session, memo_id: str) -> Optional[str]:
    """Render a memo as formatted text (delegates to format_memo_markdown)."""
    memo = db.query(LeadershipMemo).filter(LeadershipMemo.id == uuid.UUID(memo_id)).first()
    if not memo:
        return None
    return format_memo_markdown(memo)


# ---------------------------------------------------------------------------
# File persistence
# ---------------------------------------------------------------------------


def _memo_filename(week_start: datetime) -> str:
    """Return the canonical filename for a weekly memo."""
    return f"ai-platform-weekly-memo-{week_start.strftime('%Y-%m-%d')}.md"


def _docx_filename(week_start: datetime) -> str:
    """Return the canonical .docx filename for a weekly memo."""
    return f"ai-platform-weekly-memo-{week_start.strftime('%Y-%m-%d')}.docx"


def save_memo_to_file(memo: LeadershipMemo, content: str) -> Path:
    """Save the Markdown memo to the repository file system.

    Creates ``/leadership-memos/ai-platform/weekly/<filename>.md``.
    Idempotent — overwrites if file already exists for that week.
    Returns the path to the written file.
    """
    root = _repo_root()
    memo_dir = root / _MEMO_DIR
    memo_dir.mkdir(parents=True, exist_ok=True)

    filepath = memo_dir / _memo_filename(memo.week_start_date)
    filepath.write_text(content, encoding="utf-8")
    log.info("Saved memo to %s", filepath)
    return filepath


# ---------------------------------------------------------------------------
# Pandoc conversion
# ---------------------------------------------------------------------------


def convert_memo_to_docx(md_path: Path) -> Optional[Path]:
    """Convert a Markdown memo to .docx via Pandoc.

    Saves the .docx into the ``exports/`` subdirectory.
    Returns the path to the created .docx or *None* on failure.
    """
    root = _repo_root()
    export_dir = root / _EXPORT_DIR
    export_dir.mkdir(parents=True, exist_ok=True)

    docx_name = md_path.stem + ".docx"
    docx_path = export_dir / docx_name

    try:
        result = subprocess.run(
            ["pandoc", str(md_path), "-o", str(docx_path)],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode != 0:
            log.error("Pandoc conversion failed: %s", result.stderr)
            return None
        log.info("Converted memo to %s", docx_path)
        return docx_path
    except FileNotFoundError:
        log.warning("Pandoc not found — skipping .docx conversion")
        return None
    except subprocess.TimeoutExpired:
        log.error("Pandoc conversion timed out")
        return None


# ---------------------------------------------------------------------------
# Email distribution
# ---------------------------------------------------------------------------


def _build_recipient_list(db: Session) -> list[str]:
    """Return email addresses for all active platform leads.

    Only includes leads that have an email address set.
    """
    leads = _active_leads(db)
    emails: list[str] = []
    for lead in leads:
        if lead.email:
            emails.append(lead.email)
    return emails


def send_memo_email(
    db: Session,
    memo: LeadershipMemo,
    docx_path: Optional[Path] = None,
) -> bool:
    """Send the weekly memo via Resend API to the leadership team.

    Uses env vars:
      RESEND_API_KEY  — Resend API key (required)
      RESEND_FROM     — sender address (default: Signal OS <onboarding@resend.dev>)
    Recipients come from platform leads with email addresses.
    Returns True on success.
    """
    import httpx

    resend_key = os.environ.get("RESEND_API_KEY")
    if not resend_key:
        log.warning("RESEND_API_KEY not configured — skipping memo email")
        return False

    from_addr = os.environ.get("RESEND_FROM", "Signal OS <onboarding@resend.dev>")

    recipients = _build_recipient_list(db)
    if not recipients:
        log.warning("No recipients with email addresses — skipping email send")
        return False

    date_str = memo.week_start_date.strftime("%B %-d, %Y")
    subject = f"AI Platform Weekly Leadership Memo — {date_str}"

    md_content = format_memo_markdown(memo)

    body = (
        f"Please find attached the AI Platform Weekly Leadership Memo for the "
        f"week of {date_str}.\n\n"
        f"This memo covers our strategic direction, progress this week, "
        f"why it matters, next platform moves, and leadership execution.\n\n"
        f"Please review and reach out with any questions or feedback.\n\n"
        f"---\n\n{md_content}"
    )

    # Build attachments list for Resend
    attachments = []
    if docx_path and docx_path.exists():
        import base64
        with open(docx_path, "rb") as f:
            docx_b64 = base64.b64encode(f.read()).decode()
        attachments.append({
            "filename": docx_path.name,
            "content": docx_b64,
        })

    payload: dict = {
        "from": from_addr,
        "to": recipients,
        "subject": subject,
        "text": body,
    }
    if attachments:
        payload["attachments"] = attachments

    try:
        resp = httpx.post(
            "https://api.resend.com/emails",
            headers={"Authorization": f"Bearer {resend_key}"},
            json=payload,
            timeout=30,
        )
        if resp.status_code in (200, 201):
            log.info("Memo email sent to %s via Resend (id=%s)", ", ".join(recipients), resp.json().get("id"))
            return True
        else:
            log.error("Resend API error %d: %s", resp.status_code, resp.text)
            return False
    except Exception:
        log.exception("Failed to send memo email via Resend")
        return False


# ---------------------------------------------------------------------------
# Full workflow
# ---------------------------------------------------------------------------


def execute_memo_workflow(
    db: Session,
    *,
    author: Optional[str] = None,
    strategic_objective: Optional[str] = None,
    send_email: bool = True,
) -> dict:
    """Execute the full memo workflow: generate -> save -> convert -> email.

    Returns a dict with status information and file paths.
    """
    # 1. Get or generate memo for this week
    week_start = _week_start()
    existing = (
        db.query(LeadershipMemo)
        .filter(LeadershipMemo.week_start_date == week_start)
        .order_by(LeadershipMemo.created_at.desc())
        .first()
    )
    if existing:
        memo = existing
    else:
        memo = generate_memo(db, author=author, strategic_objective=strategic_objective)

    # 2. Format as Markdown
    md_content = format_memo_markdown(memo)

    # 3. Save to file
    md_path = save_memo_to_file(memo, md_content)

    # 4. Convert to .docx via Pandoc
    docx_path = convert_memo_to_docx(md_path)

    # 5. Send email
    email_sent = False
    if send_email:
        email_sent = send_memo_email(db, memo, docx_path)

    return {
        "memo_id": str(memo.id),
        "status": memo.status.value if hasattr(memo.status, "value") else memo.status,
        "md_path": str(md_path),
        "docx_path": str(docx_path) if docx_path else None,
        "email_sent": email_sent,
        "content": md_content,
    }


def get_or_generate_memo_text(db: Session, *, author: Optional[str] = None) -> str:
    """Get the latest memo for this week, or generate one. Return formatted text.

    Used by the /memo slash command.  Also triggers file persistence and
    conversion as a side-effect.
    """
    result = execute_memo_workflow(db, author=author, send_email=False)
    return result["content"]


# ---------------------------------------------------------------------------
# CRUD operations (unchanged public API)
# ---------------------------------------------------------------------------


def update_memo(
    db: Session,
    *,
    memo_id: str,
    **fields: object,
) -> Optional[LeadershipMemo]:
    """Update a memo by ID. Returns None if not found."""
    memo = db.query(LeadershipMemo).filter(LeadershipMemo.id == uuid.UUID(memo_id)).first()
    if not memo:
        return None

    for k, v in fields.items():
        if v is not None:
            if k == "status":
                v = MemoStatus(v) if isinstance(v, str) else v
            elif k in ("current_priorities", "focus_next_week", "success_criteria", "audience"):
                v = json.dumps(v) if isinstance(v, list) else v
            elif k in ("lead_updates", "dashboard_snapshot"):
                v = json.dumps(v) if isinstance(v, dict) else v
            setattr(memo, k, v)

    db.commit()
    db.refresh(memo)
    return memo


def list_memos(
    db: Session,
    *,
    status: Optional[str] = None,
    limit: int = 20,
) -> list[LeadershipMemo]:
    """List memos, newest first, optionally filtered by status."""
    q = db.query(LeadershipMemo)
    if status:
        q = q.filter(LeadershipMemo.status == MemoStatus(status))
    return q.order_by(LeadershipMemo.created_at.desc()).limit(limit).all()


def get_memo(
    db: Session,
    *,
    memo_id: str,
) -> Optional[LeadershipMemo]:
    """Get a single memo by ID."""
    return db.query(LeadershipMemo).filter(LeadershipMemo.id == uuid.UUID(memo_id)).first()


def export_memo_md(db: Session, memo_id: str) -> Optional[str]:
    """Export a memo as a markdown document. Returns None if not found."""
    memo = db.query(LeadershipMemo).filter(LeadershipMemo.id == uuid.UUID(memo_id)).first()
    if not memo:
        return None
    return format_memo_markdown(memo)


def export_memo_docx(db: Session, memo_id: str) -> Optional[bytes]:
    """Export a memo as a .docx Word document via Pandoc.

    Falls back to python-docx in-memory generation if Pandoc is not available.
    """
    memo = db.query(LeadershipMemo).filter(LeadershipMemo.id == uuid.UUID(memo_id)).first()
    if not memo:
        return None

    # Try Pandoc-based conversion first
    md_content = format_memo_markdown(memo)

    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".md", mode="w", delete=False) as tmp_md:
        tmp_md.write(md_content)
        tmp_md_path = Path(tmp_md.name)

    try:
        tmp_docx = tmp_md_path.with_suffix(".docx")
        result = subprocess.run(
            ["pandoc", str(tmp_md_path), "-o", str(tmp_docx)],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0 and tmp_docx.exists():
            docx_bytes = tmp_docx.read_bytes()
            tmp_docx.unlink(missing_ok=True)
            tmp_md_path.unlink(missing_ok=True)
            return docx_bytes
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    finally:
        tmp_md_path.unlink(missing_ok=True)

    # Fallback: python-docx in-memory generation
    return _export_docx_fallback(memo)


def _export_docx_fallback(memo: LeadershipMemo) -> bytes:
    """Generate a .docx in-memory using python-docx as a fallback."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lead_updates = _parse_json_field(memo.lead_updates, {})

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    title_para = doc.add_heading("AI Platform Weekly Leadership Memo", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Week Of: {memo.week_start_date.strftime('%B %-d, %Y')}")
    doc.add_paragraph(f"From: {memo.author or 'Leadership'}")
    status_val = memo.status.value if hasattr(memo.status, "value") else memo.status
    doc.add_paragraph(f"Status: {status_val}")

    doc.add_heading("Strategic Direction", level=1)
    doc.add_paragraph(memo.strategic_objective or DEFAULT_STRATEGIC_OBJECTIVE)

    doc.add_heading("Progress This Week", level=1)
    doc.add_paragraph(memo.progress_summary or "")

    doc.add_heading("Why It Matters", level=1)
    why_text = _parse_json_field(memo.focus_next_week, "")
    doc.add_paragraph(why_text if isinstance(why_text, str) else str(why_text))

    doc.add_heading("Next Platform Moves", level=1)
    next_moves = _parse_json_field(memo.success_criteria, "")
    doc.add_paragraph(next_moves if isinstance(next_moves, str) else str(next_moves))

    doc.add_heading("Leadership Execution", level=1)
    execution_text = ""
    if isinstance(lead_updates, dict):
        execution_text = lead_updates.pop("_execution_text", "")
    if not execution_text:
        exec_sentences: list[str] = []
        for name, info in lead_updates.items():
            if not isinstance(info, dict):
                continue
            role = info.get("role", "")
            focus = info.get("focus", "")
            exec_sentences.append(f"{name}, {role}, is focused on {focus}")
        execution_text = ". ".join(exec_sentences) + "." if exec_sentences else ""
    doc.add_paragraph(execution_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
